"""
Процессоры для извлечения текста из Word документов (.doc, .docx).
"""

import platform
import shutil
import subprocess
import tempfile
from pathlib import Path
from typing import Any, Dict, List, Optional

from ..core.base import BaseExtractor
from ..utils.logger import get_logger, log_processing_stage

logger = get_logger(__name__)


def _format_extracted_text(text: str) -> str:
    """Нормализует текст, применяя простые эвристики для заголовков."""
    if not text:
        return ""

    lines = text.split('\n')
    cleaned_lines: List[str] = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            if cleaned_lines and cleaned_lines[-1]:
                cleaned_lines.append("")
            continue

        if len(stripped) < 100 and stripped.isupper():
            cleaned_lines.append(f"# {stripped}")
        elif len(stripped) < 80 and not stripped.endswith('.'):
            cleaned_lines.append(f"## {stripped}")
        else:
            cleaned_lines.append(stripped)

    return '\n'.join(cleaned_lines)


def _extract_docx_with_python_docx(docx_path: Path) -> str:
    from docx import Document  # type: ignore

    doc = Document(str(docx_path))
    parts: List[str] = []

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_docx_with_docx2txt(docx_path: Path) -> str:
    import docx2txt  # type: ignore

    return docx2txt.process(str(docx_path))


class PythonDocxExtractor(BaseExtractor):
    """Процессор на базе python-docx для извлечения текста из .docx файлов."""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__("Python-docx", config)
        
    def supports_file_type(self, file_type: str) -> bool:
        return file_type.lower() in ['.docx']
    
    @log_processing_stage("Python-docx Extraction")
    def extract_text(self, file_path: Path, **kwargs) -> str:
        """Извлекает текст из .docx файла."""
        try:
            from docx import Document
            
            doc = Document(file_path)
            content_parts = []
            
            # Настройки извлечения
            extract_tables = self.config.get("extract_tables", True)
            extract_images = self.config.get("extract_images", False)
            preserve_formatting = self.config.get("preserve_formatting", True)
            
            # Извлекаем основной текст
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    formatted_text = self._format_paragraph(
                        paragraph, preserve_formatting
                    )
                    if formatted_text:
                        content_parts.append(formatted_text)
            
            # Извлекаем таблицы
            if extract_tables and doc.tables:
                content_parts.append("\n# Таблицы\n")
                for table_num, table in enumerate(doc.tables):
                    content_parts.append(f"## Таблица {table_num + 1}")
                    markdown_table = self._extract_table(table)
                    content_parts.append(markdown_table)
            
            # Извлекаем изображения (заглушки)
            if extract_images:
                content_parts.append("\n# Изображения\n")
                # TODO: Реализовать извлечение изображений
                content_parts.append("*Извлечение изображений пока не реализовано*")
            
            return "\n\n".join(content_parts)
            
        except ImportError:
            raise ImportError("python-docx is not installed. Run: pip install python-docx")
        except Exception as e:
            logger.error(f"Python-docx extraction failed: {e}")
            raise
    
    def _format_paragraph(self, paragraph, preserve_formatting: bool) -> str:
        """Форматирует параграф с учетом стилей."""
        if not preserve_formatting:
            return paragraph.text.strip()
        
        text = paragraph.text.strip()
        if not text:
            return ""
        
        # Определяем уровень заголовка по стилю
        style_name = paragraph.style.name.lower()
        
        if 'heading 1' in style_name or 'заголовок 1' in style_name:
            return f"# {text}"
        elif 'heading 2' in style_name or 'заголовок 2' in style_name:
            return f"## {text}"
        elif 'heading 3' in style_name or 'заголовок 3' in style_name:
            return f"### {text}"
        elif 'heading 4' in style_name or 'заголовок 4' in style_name:
            return f"#### {text}"
        elif 'heading 5' in style_name or 'заголовок 5' in style_name:
            return f"##### {text}"
        elif 'heading 6' in style_name or 'заголовок 6' in style_name:
            return f"###### {text}"
        elif 'title' in style_name or 'название' in style_name:
            return f"# {text}"
        else:
            # Проверяем форматирование текста внутри параграфа
            return self._format_runs(paragraph.runs) if paragraph.runs else text
    
    def _format_runs(self, runs) -> str:
        """Форматирует runs (части текста с разным форматированием)."""
        formatted_parts = []
        
        for run in runs:
            text = run.text
            if not text:
                continue
            
            # Применяем форматирование
            if run.bold and run.italic:
                text = f"***{text}***"
            elif run.bold:
                text = f"**{text}**"
            elif run.italic:
                text = f"*{text}*"
            
            # Можно добавить другие виды форматирования
            if hasattr(run, 'underline') and run.underline:
                text = f"<u>{text}</u>"
            
            formatted_parts.append(text)
        
        return "".join(formatted_parts)
    
    def _extract_table(self, table) -> str:
        """Извлекает таблицу в Markdown формате."""
        if not table.rows:
            return "*Пустая таблица*"
        
        markdown_lines = []
        
        # Получаем данные таблицы
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_data.append(cell_text)
            table_data.append(row_data)
        
        if not table_data:
            return "*Пустая таблица*"
        
        # Определяем максимальное количество колонок
        max_cols = max(len(row) for row in table_data)
        
        # Выравниваем все строки до максимального количества колонок
        for row in table_data:
            while len(row) < max_cols:
                row.append("")
        
        # Создаем заголовок (первая строка)
        header = "| " + " | ".join(table_data[0]) + " |"
        markdown_lines.append(header)
        
        # Разделитель
        separator = "|" + "|".join([" --- " for _ in range(max_cols)]) + "|"
        markdown_lines.append(separator)
        
        # Остальные строки
        for row in table_data[1:]:
            row_line = "| " + " | ".join(row) + " |"
            markdown_lines.append(row_line)
        
        return "\n".join(markdown_lines)


class Docx2txtExtractor(BaseExtractor):
    """Простой процессор на базе docx2txt для извлечения текста."""
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__("Docx2txt", config)
        
    def supports_file_type(self, file_type: str) -> bool:
        return file_type.lower() in ['.docx']
    
    @log_processing_stage("Docx2txt Extraction")
    def extract_text(self, file_path: Path, **kwargs) -> str:
        """Извлекает текст простым способом."""
        try:
            text = _extract_docx_with_docx2txt(file_path)

            if not text or not text.strip():
                return "Текст не извлечен или документ пуст"

            return _format_extracted_text(text)
            
        except ImportError:
            raise ImportError("docx2txt is not installed. Run: pip install docx2txt")
        except Exception as e:
            logger.error(f"Docx2txt extraction failed: {e}")
            raise


class DocExtractor(BaseExtractor):
    """
    Универсальный процессор для .doc файлов.
    Пробует несколько методов: antiword, catdoc, LibreOffice, MS Word COM.
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        super().__init__("DOC Extractor", config)
        
    def supports_file_type(self, file_type: str) -> bool:
        return file_type.lower() in ['.doc']
    
    @log_processing_stage("DOC Extraction")
    def extract_text(self, file_path: Path, **kwargs) -> str:
        """Извлекает текст из .doc файла используя доступные методы."""
        
        # Метод 1: Попытка через antiword (быстрый и надёжный)
        try:
            text = self._extract_with_antiword(file_path)
            if text.strip():
                logger.info("Successfully extracted with antiword")
                return _format_extracted_text(text)
        except Exception as e:
            logger.debug(f"antiword extraction failed: {e}")
        
        # Метод 2: Попытка через catdoc
        try:
            text = self._extract_with_catdoc(file_path)
            if text.strip():
                logger.info("Successfully extracted with catdoc")
                return _format_extracted_text(text)
        except Exception as e:
            logger.debug(f"catdoc extraction failed: {e}")
        
        # Метод 3: LibreOffice
        try:
            text = self._extract_with_libreoffice(file_path)
            if text.strip():
                logger.info("Successfully extracted with LibreOffice")
                return _format_extracted_text(text)
        except Exception as e:
            logger.debug(f"LibreOffice extraction failed: {e}")
        
        # Метод 4: MS Word COM (только Windows)
        if platform.system().lower() == "windows":
            try:
                text = self._extract_with_word_com(file_path)
                if text.strip():
                    logger.info("Successfully extracted with MS Word COM")
                    return _format_extracted_text(text)
            except Exception as e:
                logger.debug(f"MS Word COM extraction failed: {e}")
        
        raise RuntimeError(
            "Failed to extract text from .doc file. "
            "Install antiword, catdoc, LibreOffice, or MS Word for .doc support."
        )
    
    def _extract_with_antiword(self, file_path: Path) -> str:
        """Извлечение через antiword."""
        antiword_path = shutil.which('antiword')
        if not antiword_path:
            raise EnvironmentError("antiword not found")
        
        result = subprocess.run(
            [antiword_path, str(file_path)],
            capture_output=True,
            text=True,
            timeout=60,
            encoding='utf-8',
            errors='ignore'
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"antiword failed: {result.stderr}")
        
        return result.stdout
    
    def _extract_with_catdoc(self, file_path: Path) -> str:
        """Извлечение через catdoc."""
        catdoc_path = shutil.which('catdoc')
        if not catdoc_path:
            raise EnvironmentError("catdoc not found")
        
        result = subprocess.run(
            [catdoc_path, '-a', str(file_path)],
            capture_output=True,
            text=True,
            timeout=60,
            encoding='utf-8',
            errors='ignore'
        )
        
        if result.returncode != 0:
            raise RuntimeError(f"catdoc failed: {result.stderr}")
        
        return result.stdout
    
    def _extract_with_libreoffice(self, file_path: Path) -> str:
        """Извлечение через LibreOffice."""
        libreoffice_path = self._find_libreoffice()
        if not libreoffice_path:
            raise EnvironmentError("LibreOffice not found")
        
        with tempfile.TemporaryDirectory() as temp_dir:
            temp_path = Path(temp_dir)
            
            # Конвертируем в docx
            cmd = [
                libreoffice_path,
                '--headless',
                '--convert-to', 'docx:MS Word 2007 XML',
                '--outdir', str(temp_path),
                str(file_path)
            ]
            
            result = subprocess.run(
                cmd,
                capture_output=True,
                text=True,
                timeout=120
            )
            
            if result.returncode != 0:
                raise RuntimeError(f"LibreOffice conversion failed: {result.stderr}")
            
            # Ищем конвертированный файл
            converted_docx = temp_path / f"{file_path.stem}.docx"
            if not converted_docx.exists():
                candidates = list(temp_path.glob(f"{file_path.stem}*.docx"))
                if not candidates:
                    raise RuntimeError("Converted DOCX file not found")
                converted_docx = candidates[0]
            
            # Извлекаем текст из docx
            try:
                return _extract_docx_with_python_docx(converted_docx)
            except:
                return _extract_docx_with_docx2txt(converted_docx)
    
    def _extract_with_word_com(self, file_path: Path) -> str:
        """Извлечение через MS Word COM."""
        try:
            import win32com.client
            from pywintypes import com_error
        except ImportError:
            raise ImportError("pywin32 not installed")
        
        word_app = None
        doc = None
        
        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0
            
            doc = word_app.Documents.Open(
                str(file_path.absolute()),
                ReadOnly=True,
                ConfirmConversions=False,
                AddToRecentFiles=False
            )
            
            text = doc.Content.Text
            doc.Close(False)
            doc = None
            
            return text
            
        except com_error as e:
            raise RuntimeError(f"Word COM failed: {e}")
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except:
                    pass
            if word_app is not None:
                try:
                    word_app.Quit()
                except:
                    pass
    
    def _find_libreoffice(self) -> Optional[str]:
        """Ищет исполняемый файл LibreOffice."""
        possible_paths = [
            'libreoffice',
            'soffice',
            '/usr/bin/libreoffice',
            '/usr/local/bin/libreoffice',
            'C:\\Program Files\\LibreOffice\\program\\soffice.exe',
            'C:\\Program Files (x86)\\LibreOffice\\program\\soffice.exe'
        ]
        
        for path in possible_paths:
            if shutil.which(path):
                return path
            # Проверяем как прямой путь (для Windows)
            if Path(path).exists():
                return path
        
        return None


class Win32WordExtractor(BaseExtractor):
    """Процессор, использующий Microsoft Word через COM (только Windows)."""

    def __init__(self, config: Dict[str, Any] = None):
        super().__init__("MS Word COM", config)

    def supports_file_type(self, file_type: str) -> bool:
        return file_type.lower() in [".doc", ".docx", ".rtf"]

    @log_processing_stage("MS Word COM Extraction")
    def extract_text(self, file_path: Path, **kwargs) -> str:
        if platform.system().lower() != "windows":
            raise EnvironmentError("MS Word COM extractor available only on Windows")

        try:
            import win32com.client  # type: ignore
            from pywintypes import com_error  # type: ignore
        except ImportError as exc:  # pragma: no cover - depends on environment
            raise ImportError(
                "pywin32 is required for MS Word extraction. Install via 'pip install pywin32'."
            ) from exc

        word_app = None
        doc = None

        try:
            word_app = win32com.client.Dispatch("Word.Application")
            word_app.Visible = False
            word_app.DisplayAlerts = 0

            doc = word_app.Documents.Open(
                str(file_path),
                ReadOnly=True,
                ConfirmConversions=False,
                AddToRecentFiles=False,
            )

            file_format = self.config.get("file_format", 7)  # 7 = wdFormatUnicodeText
            output_encoding = self.config.get("output_encoding", "utf-16")

            with tempfile.TemporaryDirectory() as temp_dir:
                temp_path = Path(temp_dir) / f"{file_path.stem}_msword.txt"
                doc.SaveAs(str(temp_path), FileFormat=file_format)
                doc.Close(False)
                doc = None

                text = temp_path.read_text(encoding=output_encoding, errors="ignore")

            if not text.strip():
                return "Документ пуст или не удалось извлечь текст"

            # Нормализуем переносы строк
            return text.replace("\r\n", "\n").strip()

        except com_error as exc:
            raise RuntimeError(f"Microsoft Word automation failed: {exc}") from exc
        finally:
            if doc is not None:
                try:
                    doc.Close(False)
                except Exception:  # pragma: no cover - best effort cleanup
                    pass
            if word_app is not None:
                try:
                    word_app.Quit()
                except Exception:  # pragma: no cover
                    pass